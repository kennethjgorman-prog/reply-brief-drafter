"""
BriefDrafter DOCX generator: renders brief text as formatted Word documents.
"""

import re
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from werkzeug.utils import secure_filename

from src.config import PROJECTS_DIR, BRIEF_TYPE_CONFIG


def resolve_nyscef_url(page_num, nyscef_config):
    """Given a record page number, return NYSCEF URL with #page= or None."""
    if not nyscef_config or not nyscef_config.get('volumes'):
        return None
    volumes = sorted(nyscef_config['volumes'],
                     key=lambda v: v.get('first_page', 0), reverse=True)
    for vol in volumes:
        first_page = vol.get('first_page', 1)
        doc_index = vol.get('doc_index', '')
        offset = vol.get('page_offset', 0)
        if not doc_index or page_num < first_page:
            continue
        pdf_page = (page_num - first_page) + offset
        return f"https://iapps.courts.state.ny.us/nyscef/ViewDocument?docIndex={doc_index}#page={pdf_page}"
    return None


def _clean_text(text):
    """Strip markdown and fix citation formatting"""
    # Strip markdown heading markers
    text = re.sub(r'^#{1,6}\s*', '', text)
    # Convert **bold** to plain text (bold handled separately for headings)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = text.replace('**', '')
    # Remove stray asterisks used for emphasis
    text = re.sub(r'(?<!\w)\*([^*]+)\*(?!\w)', r'\1', text)
    # Fix record citation prefixes
    text = re.sub(r'\(R\.\s*(\d+[^)]*)\)', r'(\1)', text)
    text = re.sub(r'\(A\.\s*(\d+[^)]*)\)', r'(\1)', text)
    # Fix case citation periods: A.D.3d -> AD3d, A.D.2d -> AD2d, N.Y.S.2d -> NYS2d, etc.
    text = re.sub(r'A\.D\.3d', 'AD3d', text)
    text = re.sub(r'A\.D\.2d', 'AD2d', text)
    text = re.sub(r'N\.Y\.S\.3d', 'NYS3d', text)
    text = re.sub(r'N\.Y\.S\.2d', 'NYS2d', text)
    text = re.sub(r'N\.Y\.3d', 'NY3d', text)
    text = re.sub(r'N\.Y\.2d', 'NY2d', text)
    text = re.sub(r'N\.E\.3d', 'NE3d', text)
    text = re.sub(r'N\.E\.2d', 'NE2d', text)
    text = re.sub(r'Misc\.?\s*3d', 'Misc 3d', text)
    text = re.sub(r'Misc\.?\s*2d', 'Misc 2d', text)
    return text


def _is_short_heading(text):
    """Short headings like POINT I, PRELIMINARY STATEMENT, CONCLUSION — stay double-spaced centered."""
    clean = text.strip()
    return bool(re.match(r'^POINT\s+[IVXLCDM\d]+:?\s*$', clean) or
                clean in ('PRELIMINARY STATEMENT', 'CONCLUSION', 'ARGUMENT',
                          'STATEMENT OF THE CASE', 'STATEMENT OF FACTS',
                          'QUESTIONS PRESENTED', 'DISCUSSION', 'INTRODUCTION',
                          'COUNTERSTATEMENT OF FACTS'))


def _is_heading(text):
    """Detect if a line is a section heading (ALL CAPS, short, no period at end)"""
    stripped = text.strip()
    if not stripped:
        return False
    # Remove leading tabs for detection
    clean = stripped.lstrip('\t').strip()
    if not clean:
        return False
    # Known heading patterns
    heading_patterns = [
        r'^POINT\s+[IVXLCDM\d]+:?',
        r'^PRELIMINARY STATEMENT',
        r'^STATEMENT OF THE CASE',
        r'^STATEMENT OF FACTS',
        r'^COUNTERSTATEMENT',
        r'^QUESTIONS PRESENTED',
        r'^ARGUMENT',
        r'^CONCLUSION',
        r'^DISCUSSION',
        r'^BRIEF FOR',
        r'^REPLY BRIEF',
        r'^SUPREME COURT',
        r'^APPELLATE DIVISION',
    ]
    for pattern in heading_patterns:
        if re.match(pattern, clean):
            return True
    # General ALL CAPS detection: mostly uppercase letters, no lowercase sentences
    alpha_chars = [c for c in clean if c.isalpha()]
    if len(alpha_chars) > 3:
        upper_ratio = sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars)
        if upper_ratio > 0.85 and len(clean) < 500 and len(clean.split()) >= 4:
            # Skip if it looks like a party name (corporate suffixes)
            if not re.search(r'\b(?:Corp|Inc|LLC|Ltd)\b', clean):
                return True
    return False


def _is_subheading(text):
    """Detect sub-headings like A., B., 1., 2. at start"""
    stripped = text.strip().lstrip('\t')
    if len(stripped) > 150 or stripped.endswith('.'):
        return False
    return bool(re.match(r'^[A-Z]\.\s', stripped) or re.match(r'^\d+\.\s', stripped))


def _add_hyperlink(paragraph, url, text, font_size_pt=12):
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Courier New')
    rFonts.set(qn('w:hAnsi'), 'Courier New')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size_pt * 2))
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    run_elem.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _add_run(p, text, is_bold=False):
    """Add a plain Courier New 12pt run."""
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(12)
    if is_bold:
        run.bold = True


def _add_text_with_citations(p, text, nyscef_cfg, is_bold=False):
    """Split text on record citations, inserting hyperlinks where NYSCEF URLs resolve.
    Handles comma-separated refs like (547-548, 556) and (730-731, 734) and single-digit (4)."""
    # Match bare page-number citations: (4), (5-6), (47, 55), (547-548, 556)
    # Exclude: court citations (2d Dept 2020), years standing alone (2023-2026)
    citation_pat = re.compile(r'(?<![a-zA-Z0-9\u00a7\u00b6)\-])(\(\d+(?:\s*-\s*\d+)?(?:,\s*\d+(?:\s*-\s*\d+)?)*\))(?![a-zA-Z])')
    segments = citation_pat.split(text)
    for segment in segments:
        if segment.startswith('(') and segment.endswith(')') and re.match(r'\(\d', segment):
            inner = segment[1:-1]
            parts = [pt.strip() for pt in inner.split(',')]
            _add_run(p, '(', is_bold)
            for i, part in enumerate(parts):
                if i > 0:
                    _add_run(p, ', ', is_bold)
                page = int(re.match(r'(\d+)', part).group(1))
                url = resolve_nyscef_url(page, nyscef_cfg)
                if url:
                    _add_hyperlink(p, url, part)
                else:
                    _add_run(p, part, is_bold)
            _add_run(p, ')', is_bold)
        elif segment:
            _add_run(p, segment, is_bold)


def _add_paragraph(doc, text, nyscef_cfg=None, is_bold=False, alignment=None,
                    link_citations=False, is_subheading=False, single_spaced=False):
    """Add a paragraph with Courier New 12pt, formatted per Rosman exemplar specs"""
    text = _clean_text(text)
    p = doc.add_paragraph()
    if is_subheading:
        # Subheadings: justified, single-spaced, all lines indented 0.5"
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif single_spaced:
        # Long point headings: centered, single-spaced, first-line indent 0.5"
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Inches(0.5)
    else:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    if alignment:
        p.alignment = alignment
    elif not is_bold:
        # Body paragraphs: justified with first-line indent
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if not text.startswith('\t'):
            p.paragraph_format.first_line_indent = Inches(0.5)

    use_links = link_citations and nyscef_cfg

    # Split on underscored case names
    parts = re.split(r'(_[^_]+_)', text)
    for part in parts:
        if part.startswith('_') and part.endswith('_') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(12)
            run.underline = True
            if is_bold:
                run.bold = True
        elif use_links:
            _add_text_with_citations(p, part, nyscef_cfg, is_bold)
        else:
            _add_run(p, part, is_bold)
    return p


def generate_brief_docx(project):
    """Generate complete brief as Word document (type-aware).

    Args:
        project: Project dict (must include 'id' key).

    Returns:
        Path to the saved .docx file.
    """
    project_id = project['id']
    brief_type = project.get('brief_type', 'reply')
    config = BRIEF_TYPE_CONFIG.get(brief_type, BRIEF_TYPE_CONFIG['reply'])

    # Create Word document
    doc = DocxDocument()

    # Set 1-inch margins on all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default style to Courier New, 12pt, double-spaced
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    nyscef_cfg = project.get('nyscef_config')

    def add_para(doc, text, is_bold=False, alignment=None, link_citations=False,
                 is_subheading=False, single_spaced=False):
        return _add_paragraph(doc, text, nyscef_cfg=nyscef_cfg, is_bold=is_bold,
                              alignment=alignment, link_citations=link_citations,
                              is_subheading=is_subheading, single_spaced=single_spaced)

    def add_spacer(doc):
        """Add a single-spaced blank paragraph for spacing between headings."""
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # --- Build the document ---

    # Appellate caption block — single-spaced, left-aligned
    appellant = project.get('appellant', '').strip()
    respondent = project.get('respondent', '').strip()
    docket_number = project.get('docket_number', '').strip()
    brief_type = project.get('brief_type', 'appellant')
    representing = project.get('representing', 'appellant')
    attorney_name = project.get('attorney_name', '').strip()
    attorney_firm = project.get('attorney_firm', '').strip()

    court_name = project.get('court', '').strip()
    if not court_name or 'appellate' not in court_name.lower():
        court_name = 'APPELLATE DIVISION, SECOND DEPARTMENT'

    # Determine party designations
    if representing == 'appellant':
        appellant_designation = 'Plaintiff-Appellant'
        respondent_designation = 'Defendants-Respondents'
    else:
        appellant_designation = 'Plaintiff-Respondent'
        respondent_designation = 'Defendants-Appellants'

    # Determine brief title label (short form for caption)
    brief_title_map = {
        'appellant': 'BRIEF FOR APPELLANT',
        'respondent': 'BRIEF FOR RESPONDENT',
        'reply': 'REPLY BRIEF',
    }
    caption_title = brief_title_map.get(brief_type, config['doc_title'])

    def _caption_para(text, bold=False, right_text=None):
        """Add a single-spaced caption paragraph, left-aligned."""
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if right_text:
            # Left text + right-aligned text on same line using tab stop
            tab_stop = OxmlElement('w:tab')
            pPr = p._p.get_or_add_pPr()
            tabs = OxmlElement('w:tabs')
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), '9360')  # right margin at 6.5"
            tabs.append(tab)
            pPr.append(tabs)
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(12)
            run.bold = bold
            run2 = p.add_run('\t')
            run2.font.name = 'Courier New'
            run2.font.size = Pt(12)
            run3 = p.add_run(right_text)
            run3.font.name = 'Courier New'
            run3.font.size = Pt(12)
            run3.bold = True
        else:
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(12)
            run.bold = bold
        return p

    # Court heading — NY appellate courts need full two-line header
    upper_court = court_name.upper()
    if 'APPELLATE DIVISION' in upper_court:
        _caption_para('SUPREME COURT OF THE STATE OF NEW YORK', bold=True)
        _caption_para(upper_court, bold=True)
    else:
        _caption_para(upper_court, bold=True)

    # Dashed line
    _caption_para('--------------------------------------X')

    # Appellant name with brief title right-aligned
    _caption_para(appellant.upper(), right_text=caption_title)

    # Party designation indented
    _caption_para(f'\t{appellant_designation}')

    # -against-
    _caption_para('')
    _caption_para('\t-against-')
    _caption_para('')

    # Respondent name
    _caption_para(respondent.upper() + ',')

    # Respondent designation indented
    _caption_para(f'\t{respondent_designation}')

    # Check if case_name has additional parties (e.g., "and THE COUNTY OF NASSAU")
    case_name = project.get('case_name', '')
    # Look for parties after the respondent
    resp_upper = respondent.upper()
    if resp_upper in case_name.upper():
        remainder = case_name.upper()[case_name.upper().index(resp_upper) + len(resp_upper):].strip()
        if remainder.startswith('AND '):
            additional_party = remainder[4:].strip()
            _caption_para('')
            _caption_para('\t-and-')
            _caption_para('')
            _caption_para(additional_party)
            _caption_para('\tDefendant')

    # Closing dashed line
    _caption_para('--------------------------------------X')

    # Blank line to transition to double-spaced body
    add_para(doc, "")

    # Add drafted sections
    sections = project.get('drafted_sections', {})

    def _strip_ai_caption(text):
        """Strip AI-generated caption/title block from the top of drafted content.
        The AI often generates its own court heading, party block, and brief title
        which duplicates the DOCX generator's caption."""
        lines = text.split('\n')
        # Find where the actual brief content starts (PRELIMINARY STATEMENT, POINT I, etc.)
        caption_keywords = [
            'SUPREME COURT', 'APPELLATE DIVISION', 'Plaintiff-Appellant',
            'Plaintiff-Respondent', 'Defendants-Respondents', 'Defendants-Appellants',
            '-against-', 'REPLY BRIEF FOR', 'BRIEF FOR APPELLANT',
            'BRIEF FOR RESPONDENT', 'BRIEF FOR PLAINTIFF', 'Docket No.',
            'Index No.',
        ]
        content_starts = [
            'PRELIMINARY STATEMENT', 'INTRODUCTION', 'STATEMENT OF THE CASE',
            'STATEMENT OF FACTS', 'POINT I', 'QUESTIONS PRESENTED',
            'TABLE OF CONTENTS', 'TABLE OF AUTHORITIES',
        ]
        first_content_idx = None
        for i, line in enumerate(lines):
            stripped = line.strip().upper()
            for cs in content_starts:
                if stripped.startswith(cs):
                    first_content_idx = i
                    break
            if first_content_idx is not None:
                break

        if first_content_idx and first_content_idx > 0:
            # Check if the lines before first_content_idx look like a caption
            pre_lines = '\n'.join(lines[:first_content_idx]).upper()
            is_caption = any(kw.upper() in pre_lines for kw in caption_keywords)
            if is_caption:
                return '\n'.join(lines[first_content_idx:])
        return text

    # Prefer individual argument sections over full_brief when argument sections exist
    has_individual_args = any(k.startswith('argument_') for k in sections)
    if 'full_brief' in sections and not has_individual_args:
        content = _strip_ai_caption(sections['full_brief'].get('content', ''))
        # Strip validation artifacts
        content = re.sub(r'\s*\[CITE NEEDED\]\.?', '.', content)
        content = re.sub(r'\s*\[CASE CITE NEEDED\]\.?', '.', content)
        content = re.sub(r'\s*\[FULL CITE NEEDED\]\.?', '.', content)
        content = re.sub(r'\s*\[UNVERIFIED CITATION\]\.?', '', content)
        content = re.sub(r'\s*\[CITE NUMBER UNVERIFIED\]\.?', '', content)
        content = re.sub(r'\s*\[VERIFY\]\.?', '.', content)
        content = re.sub(r'\.\.', '.', content)

        prev_type = None
        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            elif _is_heading(line):
                if _is_short_heading(stripped):
                    add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    add_spacer(doc)
                else:
                    add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, single_spaced=True)
                    add_spacer(doc)
                prev_type = 'heading'
            elif _is_subheading(line):
                add_para(doc, line.strip(), is_bold=True, is_subheading=True, link_citations=True)
                add_spacer(doc)
                prev_type = 'subheading'
            else:
                add_para(doc, line, link_citations=True)
                prev_type = 'body'

    else:
        if 'facts' in sections:
            for line in sections['facts'].get('content', '').split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                elif _is_heading(line):
                    if _is_short_heading(stripped):
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    else:
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, single_spaced=True)
                elif _is_subheading(line):
                    add_para(doc, line.strip(), is_bold=True, is_subheading=True, link_citations=True)
                else:
                    add_para(doc, line, link_citations=True)

        if 'procedural_history' in sections:
            for line in sections['procedural_history'].get('content', '').split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                elif _is_heading(line):
                    if _is_short_heading(stripped):
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    else:
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, single_spaced=True)
                elif _is_subheading(line):
                    add_para(doc, line.strip(), is_bold=True, is_subheading=True, link_citations=True)
                else:
                    add_para(doc, line, link_citations=True)

        if 'intro' in sections:
            add_para(doc, "PRELIMINARY STATEMENT", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for line in sections['intro'].get('content', '').split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                elif _is_heading(line):
                    if _is_short_heading(stripped):
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    else:
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, single_spaced=True)
                elif _is_subheading(line):
                    add_para(doc, line.strip(), is_bold=True, is_subheading=True, link_citations=True)
                else:
                    add_para(doc, line, link_citations=True)

        arg_sections = [(k, v) for k, v in sections.items() if k.startswith('argument_')]
        arg_sections.sort(key=lambda x: int(x[0].split('_')[1]) if x[0].split('_')[1].isdigit() else 0)

        # Warn about missing argument sections if brief_structure defines expected points
        brief_structure = project.get('brief_structure')
        if brief_structure and 'points' in brief_structure:
            for point in brief_structure['points']:
                arg_key = f"argument_{point['id']}"
                if arg_key not in sections:
                    add_para(doc, f"[POINT {point['id']} NOT YET DRAFTED]", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        for i, (key, section) in enumerate(arg_sections, 1):
            add_para(doc, f"POINT {i}", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for line in section.get('content', '').split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                elif _is_heading(line):
                    if _is_short_heading(stripped):
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    else:
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, single_spaced=True)
                elif _is_subheading(line):
                    add_para(doc, line.strip(), is_bold=True, is_subheading=True, link_citations=True)
                else:
                    add_para(doc, line, link_citations=True)
                    prev_type = 'body'

        # Output any custom sections not already handled
        known_keys = {'facts', 'procedural_history', 'intro', 'conclusion', 'full_brief'}
        known_keys.update(k for k in sections if k.startswith('argument_'))
        for key in sections:
            if key in known_keys:
                continue
            section = sections[key]
            content = section.get('content', '')
            if not content:
                continue
            for line in content.split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                elif _is_heading(line):
                    if _is_short_heading(stripped):
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    else:
                        add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, single_spaced=True)
                elif _is_subheading(line):
                    add_para(doc, line.strip(), is_bold=True, is_subheading=True, link_citations=True)
                else:
                    add_para(doc, line, link_citations=True)

        if 'conclusion' in sections:
            add_para(doc, "CONCLUSION", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for line in sections['conclusion'].get('content', '').split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                elif _is_heading(line):
                    add_para(doc, stripped, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                elif _is_subheading(line):
                    add_para(doc, line.strip(), is_bold=True, is_subheading=True, link_citations=True)
                else:
                    add_para(doc, line, link_citations=True)

    # Dynamic signature block
    add_para(doc, "")
    add_para(doc, "Respectfully submitted,")
    add_para(doc, "")
    add_para(doc, "_______________________")
    add_para(doc, project.get('attorney_name', ''))
    add_para(doc, project.get('attorney_firm', ''))
    add_para(doc, config['signature_role'])

    # Save with dynamic filename
    output_filename = config['output_filename']
    output_path = PROJECTS_DIR / project_id / output_filename
    doc.save(output_path)

    return output_path


def generate_section_docx(project, section_key):
    """Download a single drafted section as a Word document.

    Args:
        project: Project dict (must include 'id' key).
        section_key: Key into project['drafted_sections'].

    Returns:
        Tuple of (output_path, filename).

    Raises:
        KeyError: If section_key is not found or has no content.
    """
    project_id = project['id']
    sections = project.get('drafted_sections', {})
    if section_key not in sections or not sections[section_key].get('content'):
        raise KeyError(f"Section '{section_key}' not found or empty")

    content = sections[section_key]['content']
    case_name = project.get('case_name', 'draft')

    nyscef_cfg = project.get('nyscef_config')

    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    nyscef_cfg = project.get('nyscef_config')

    # Set 1-inch margins
    for section_obj in doc.sections:
        section_obj.top_margin = Inches(1)
        section_obj.bottom_margin = Inches(1)
        section_obj.left_margin = Inches(1)
        section_obj.right_margin = Inches(1)

    # Use the same _add_paragraph as the main generator
    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if _is_heading(line):
            if _is_short_heading(stripped):
                _add_paragraph(doc, stripped, nyscef_cfg=nyscef_cfg, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                _add_paragraph(doc, stripped, nyscef_cfg=nyscef_cfg, is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, single_spaced=True)
        elif _is_subheading(line):
            _add_paragraph(doc, stripped, nyscef_cfg=nyscef_cfg, is_bold=True, is_subheading=True, link_citations=True)
        else:
            _add_paragraph(doc, line, nyscef_cfg=nyscef_cfg, link_citations=True)

    label = section_key.replace('_', ' ').title()
    case_safe = secure_filename(case_name)
    filename = f"{label}_{case_safe}.docx"

    output_path = PROJECTS_DIR / project_id / filename
    doc.save(output_path)

    return output_path, filename
